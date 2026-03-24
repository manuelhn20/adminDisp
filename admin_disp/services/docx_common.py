"""docx_common.py

Utilidades compartidas para generacion de documentos (DOCX/PDF).

Centraliza:
  - Formateo de numero de linea y costo con moneda
  - Reemplazo robusto de placeholders en documentos DOCX (incluye [LOGO])

Notas:
  - Este modulo NO depende de Flask.
  - Los placeholders en plantillas usan el formato: [CLAVE] (siempre MAYUSCULAS).
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Optional

from docx.document import Document as DocxDocument
from docx.shared import Mm

logger = logging.getLogger("admin_disp.services.docx_common")

# ---------------------------------------------------------------------------
# Formateo
# ---------------------------------------------------------------------------

def format_numero_linea(numero_linea_raw: Optional[str]) -> str:
    """Formatea numero de linea al formato +(COD) XXXX-XXXX.

    Acepta entradas como '5038765432', '503-8765432', '503 8765432'.
    Retorna el valor original si no puede inferir el patron.
    """
    if not numero_linea_raw:
        return "N/A"

    raw = str(numero_linea_raw).strip()
    if not raw:
        return "N/A"

    # Caso con espacio: "503 8765432"
    parts = raw.split()
    if len(parts) == 2 and all(p.isdigit() for p in parts):
        ext, num = parts
        return f"+({ext}) {num[:4]}-{num[4:]}" if len(num) >= 4 else f"+({ext}) {num}"

    digits = re.sub(r"\D", "", raw)

    if len(digits) == 10:
        return f"+({digits[:3]}) {digits[3:7]}-{digits[7:]}"

    for ext_len in (3, 2, 1):
        if len(digits) > ext_len + 4:
            ext, num = digits[:ext_len], digits[ext_len:]
            return f"+({ext}) {num[:4]}-{num[4:]}"

    return raw


def format_costo_con_moneda(costo_plan: Any, moneda_plan: Optional[str]) -> str:
    """Formatea costo con simbolo de moneda.

    Ejemplos:
      (27.50, '$')  -> '$27.50'
      (525.50, 'L') -> 'L525.50'
      (100, 'USD')  -> '$100.00'
    """
    if costo_plan is None or costo_plan == "":
        return "N/A"

    try:
        costo_float = float(costo_plan)
    except (ValueError, TypeError) as e:
        logger.debug("No se pudo convertir costo_plan '%s' a float: %s", costo_plan, e)
        return "N/A"

    moneda = (moneda_plan or "L").strip().upper()
    simbolo = "$" if moneda in ("USD", "$") else ("L" if moneda == "L" else moneda)
    return f"{simbolo}{costo_float:.2f}"


# ---------------------------------------------------------------------------
# Opciones de reemplazo
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class PlaceholderOptions:
    """Opciones para el motor de reemplazo de placeholders."""

    logo_path: Optional[Path] = None
    logo_width_mm: float = 25.0


# ---------------------------------------------------------------------------
# Logica interna de reemplazo
# ---------------------------------------------------------------------------

def _insert_logo(paragraph, runs: list, runs_text: list, logo_path: Path, width_mm: float) -> None:
    """Reemplaza [LOGO] en el parrafo por la imagen del logo."""
    placeholder = "[LOGO]"
    full_text = "".join(runs_text)
    idx = full_text.find(placeholder)
    if idx == -1:
        return
    if not logo_path.exists():
        logger.warning("Imagen de logo no encontrada: %s", logo_path)
        return

    end_idx = idx + len(placeholder)
    cum = 0
    s_run = s_off = 0
    for i, t in enumerate(runs_text):
        if cum + len(t) > idx:
            s_run, s_off = i, idx - cum
            break
        cum += len(t)

    cum2, e_run, e_off = cum, s_run, 0
    for j in range(s_run, len(runs_text)):
        if cum2 + len(runs_text[j]) >= end_idx:
            e_run, e_off = j, end_idx - cum2
            break
        cum2 += len(runs_text[j])

    prefix = runs_text[s_run][:s_off]
    suffix = runs_text[e_run][e_off:]

    if prefix:
        paragraph.insert_paragraph_before().add_run(prefix)

    runs[s_run].text = ""
    runs[s_run].add_picture(str(logo_path), width=Mm(width_mm))

    if suffix:
        paragraph.add_run(suffix)

    for k in range(s_run + 1, e_run + 1):
        runs[k].text = ""


def _replace_text_in_paragraph(paragraph, replacements: Dict[str, Any]) -> None:
    """Reemplaza placeholders [CLAVE] preservando el máximo formato posible.
    
    Estrategia:
      1. Si placeholder está en UN SOLO run: reemplazo directo (preserva formato)
      2. Si placeholder se divide entre múltiples runs:
         - Coloca valor en primer run (respeta su formato/style)
         - Vacía runs intermedios sin eliminarlos (preservan estructura XML)
         - Mantiene suffix en último run
    
    Nota: La búsqueda es case-insensitive para [MESES], [meses], [Meses], etc.
    """
    if not paragraph.runs:
        return

    runs = list(paragraph.runs)
    runs_text = [r.text or "" for r in runs]
    full_text = "".join(runs_text)

    for key, value in replacements.items():
        if value is None or value == "":
            continue
        
        # Buscar variación case-insensitive del placeholder
        # Ej: [MESES], [meses], [Meses], etc.
        placeholder_upper = f"[{key.upper()}]"
        str_value = str(value)
        search_from = 0
        
        # DEBUG: Log para detectar placeholders no reemplazados
        found_count = 0

        while True:
            # Buscar el placeholder de forma case-insensitive
            idx = full_text.upper().find(placeholder_upper, search_from)
            if idx == -1:
                if found_count == 0 and len(full_text) > 0:
                    # Placeholder no encontrado en este párrafo
                    logger.debug(f"Placeholder [{key}] no encontrado en párrafo (se buscaba {placeholder_upper}). Texto: '{full_text[:100]}...'")
                break
            
            found_count += 1
            logger.debug(f"Encontrado placeholder [{key}] ocurrencia #{found_count} en posición {idx}")
            
            end_idx = idx + len(placeholder_upper)

            # Encontrar run inicial (donde comienza el placeholder)
            cum = 0
            s_run = s_off = 0
            for i, t in enumerate(runs_text):
                if cum + len(t) > idx:
                    s_run, s_off = i, idx - cum
                    break
                cum += len(t)

            # Encontrar run final (donde termina el placeholder)
            cum2, e_run, e_off = cum, s_run, 0
            for j in range(s_run, len(runs_text)):
                if cum2 + len(runs_text[j]) >= end_idx:
                    e_run, e_off = j, end_idx - cum2
                    break
                cum2 += len(runs_text[j])

            prefix = runs_text[s_run][:s_off]
            suffix = runs_text[e_run][e_off:]

            if s_run == e_run:
                # ✓ Caso simple: placeholder completamente en un solo run
                #   Preserva 100% del formato del run original
                runs[s_run].text = prefix + str_value + suffix
                runs_text[s_run] = prefix + str_value + suffix
                logger.debug(f"  → Reemplazado en run único {s_run}: [{key}] = '{str_value}'")
            else:
                # Caso complejo: placeholder dividido entre múltiples runs
                runs[s_run].text = prefix + str_value
                runs_text[s_run] = prefix + str_value
                
                # Vaciar runs intermedios
                for k in range(s_run + 1, e_run):
                    runs[k].text = ""
                    runs_text[k] = ""
                
                # Asignar suffix en último run
                runs[e_run].text = suffix
                runs_text[e_run] = suffix
                logger.debug(f"  → Reemplazado en runs {s_run}-{e_run}: [{key}] = '{str_value}'")

            # Actualizar búsqueda para siguiente ocurrencia
            full_text = "".join(runs_text)
            search_from = idx + len(str_value)

    # Eliminar solo parrafos totalmente vacios tras los reemplazos
    para_text = "".join(r.text or "" for r in paragraph.runs).replace("\u200B", "").strip()
    if not para_text:
        try:
            para_elem = paragraph._element
            para_parent = para_elem.getparent()
            if para_parent is not None:
                para_parent.remove(para_elem)
        except Exception:
            for r in paragraph.runs:
                r.text = ""


# ---------------------------------------------------------------------------
# API publica principal
# ---------------------------------------------------------------------------

def replace_placeholders(
    doc: DocxDocument,
    replacements: Dict[str, Any],
    *,
    options: Optional[PlaceholderOptions] = None,
) -> None:
    """Reemplaza placeholders [CLAVE] en todo el documento DOCX.

    Los keys en `replacements` deben ir en MAYUSCULAS para coincidir con las
    plantillas (ej. {'NOMBRE_EMPLEADO': 'Juan'} reemplaza [NOMBRE_EMPLEADO]).

    Soporta:
      - Texto:  [CLAVE]
      - Logo:   [LOGO]  (requiere options.logo_path valido)
    """
    opts = options or PlaceholderOptions()
    # Normalizar keys a mayusculas para que el motor sea case-insensitive
    effective_repl = {k.upper(): v for k, v in (replacements or {}).items()}

    def _process_paragraph(p):
        if not p.runs:
            return
        runs = list(p.runs)
        runs_text = [r.text or "" for r in runs]
        full_text = "".join(runs_text)
        if opts.logo_path and "[LOGO]" in full_text:
            _insert_logo(p, runs, runs_text, opts.logo_path, opts.logo_width_mm)
        _replace_text_in_paragraph(p, effective_repl)

    def _process_table(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _process_paragraph(p)

    for p in doc.paragraphs:
        _process_paragraph(p)
    for t in doc.tables:
        _process_table(t)

    for section in doc.sections:
        for p in section.header.paragraphs:
            _process_paragraph(p)
        for t in section.header.tables:
            _process_table(t)
        for p in section.footer.paragraphs:
            _process_paragraph(p)
        for t in section.footer.tables:
            _process_table(t)


# ---------------------------------------------------------------------------
# RemapData  (compatibilidad con pdf_service y logica legacy)
# ---------------------------------------------------------------------------

class RemapData:
    """Funciones de remapeo para plantillas con tokens de formato antiguo."""

    @staticmethod
    def replace_text_in_paragraphs(doc, replacements: dict) -> None:
        """Reemplaza texto en parrafos preservando formato de runs."""
        for paragraph in doc.paragraphs:
            full_text = "".join(r.text for r in paragraph.runs)
            if not any(k in full_text for k in replacements):
                continue
            for old, new in replacements.items():
                if old not in full_text:
                    continue
                for run in paragraph.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                full_text = "".join(r.text for r in paragraph.runs)

    @staticmethod
    def replace_text_in_tables(doc, replacements: dict) -> None:
        """Reemplaza texto en tablas preservando formato de runs."""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = "".join(r.text for r in paragraph.runs)
                        if not any(k in full_text for k in replacements):
                            continue
                        for old, new in replacements.items():
                            if old not in full_text:
                                continue
                            for run in paragraph.runs:
                                if old in run.text:
                                    run.text = run.text.replace(old, new)
                            full_text = "".join(r.text for r in paragraph.runs)

    @staticmethod
    def load_template_mapping(template_name: str) -> dict:
        """Retorna mapping predefinido para plantillas conocidas."""
        if template_name == "celular":
            return {
                "MOTOROLA": "{{Marca}}",
                "{{MOTOROLA}}": "{{Marca}}",
                "E13": "{{Modelo}}",
                "{{E13}}": "{{Modelo}}",
                "$21.99": "{{Costo}}",
                "11/11/2025": "{{Fecha_asignacion}}",
                "92767760": "{{Linea}}",
                "3511651485666": "{{IMEI}}",
                "NOMBRE_EMPLEADO": "{{Nombre_empleado}}",
                "{{NOMBRE_EMPLEADO}}": "{{Nombre_empleado}}",
                "{{i6s123}}": "{{IMEI}}",
                "i6s123": "{{IMEI}}",
            }
        return {}

    @staticmethod
    def apply_template_mapping(doc, template_map: dict, common_replacements: dict, expanded_replacements: dict) -> None:
        """Traduce tokens legacy a sus valores finales y aplica reemplazos."""
        resolved: dict = {}
        all_repl = {**common_replacements, **expanded_replacements}

        for orig_token, target_token in template_map.items():
            final_val = None
            if isinstance(target_token, str):
                # Buscar directamente
                if target_token in all_repl:
                    final_val = all_repl[target_token]
                else:
                    # Intentar variantes del token {{Clave}}
                    inner = target_token[2:-2] if target_token.startswith("{{") and target_token.endswith("}}") else target_token
                    for cand in (target_token, inner, inner.upper(), inner.lower(),
                                 inner.replace("_", " "), inner.replace(" ", "_")):
                        if cand in all_repl:
                            final_val = all_repl[cand]
                            break
            resolved[orig_token] = final_val if final_val is not None else (target_token if isinstance(target_token, str) else "")

        if resolved:
            RemapData.replace_text_in_paragraphs(doc, resolved)
            RemapData.replace_text_in_tables(doc, resolved)
