from __future__ import annotations

"""docexp.py

Refactorización:
- Helpers `_iter_all_paragraphs` y `_iter_all_paragraphs_except_signatures` eliminan
  la repetición del patrón paragraphs / tablas / header / footer que aparecía 6+ veces.
- `_render_docx_to_bytesio` centraliza la lógica de serialización DOCX (antes duplicada
  en `_render_docx_bytes` y `_render_docx_bytes_without_signatures`).
- `_save_tmp_docx` centraliza la escritura en disco + validación ZIP.
- `_get_correlativo` centraliza la obtención del correlativo para cada prefix.
- `_export_single` elimina la triplicación de export_tablet / export_laptop / export_periferico.
- Los endpoints Flask de /*/generate se unifican en `_generate_endpoint`.
- Los endpoints Flask de /*/fields se unifican en `_fields_endpoint`.
"""

import base64
import gc
import logging
import os
import tempfile
import time
import uuid
import zipfile
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Optional, Set, Tuple

from docx import Document
from docx.shared import Inches, Mm

try:
    from .docx_common import replace_placeholders as _replace_placeholders, PlaceholderOptions
except Exception:
    from docx_common import replace_placeholders as _replace_placeholders, PlaceholderOptions

try:
    from docx2pdf import convert as _docx2pdf_convert
    HAS_DOCX2PDF = True
except Exception as _e:
    _docx2pdf_convert = None
    HAS_DOCX2PDF = False


# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
_LOG_DIR = Path(__file__).parent.parent.parent / "logs"
_LOG_DIR.mkdir(exist_ok=True)

logger = logging.getLogger("admin_disp.services.docexp")
logger.setLevel(logging.DEBUG)

if not logger.handlers:
    _handler = logging.FileHandler(str(_LOG_DIR / "docexp.log"), encoding="utf-8")
    _handler.setFormatter(logging.Formatter(
        "[%(asctime)s] %(levelname)s - %(name)s - %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    ))
    logger.addHandler(_handler)

if HAS_DOCX2PDF:
    logger.info("docx2pdf importado exitosamente")
else:
    logger.warning("docx2pdf no disponible")


# ---------------------------------------------------------------------------
# Templates & field sets
# ---------------------------------------------------------------------------
_FORMS_DIR = Path(__file__).parent.parent / "form"

def _resolve_template_path(*candidate_names: str) -> Path:
    """Retorna el primer template existente para mantener compatibilidad entre nombres."""
    for name in candidate_names:
        candidate = _FORMS_DIR / name
        if candidate.exists():
            return candidate
    return _FORMS_DIR / candidate_names[0]


TEMPLATE_PRO_TI_001 = _resolve_template_path(
    "PRO-TI-CE-001-CORRELATIVO CERTIFICADO DE COMPROMISO Y ENTREGA DE TELEFONO CORPORATIVO.docx",
)
TEMPLATE_PRO_TI_002 = _resolve_template_path(
    "PRO-TI-CE-002-CORRELATIVO MEMORANDO DE ENTREGA.docx",
)
TEMPLATE_PRO_TI_003 = _resolve_template_path(
    "PRO-TI-CE-003-CORRELATIVO CERTIFICADO DE COMPROMISO Y ENTREGA DE TABLET.docx",
    "PRO-TI-CE-003-CORRELATIVO ENTREGA DE TABLET.docx",
)
TEMPLATE_PRO_TI_004 = _resolve_template_path(
    "PRO-TI-CE-004-CORRELATIVO CERTIFICADO ENTREGA DE COMPUTADORA.docx",
)
TEMPLATE_PRO_TI_005 = _resolve_template_path(
    "PRO-TI-CE-005-CORRELATIVO ENTREGA DE PERIFERICO.docx",
)

TIPOS_TELEFONO: Set[str] = {"Celular"}
TIPOS_PC: Set[str] = {"Laptop"}
TIPOS_TABLET: Set[str] = {"Tablet"}
TIPOS_PERIFERICOS: Set[str] = {"Teclado", "Mouse", "Auriculares", "Monitor", "Impresora", "Teléfono VoIP", "Router", "Switch", "Adaptador"}

TELEFONO_FIELDS: Set[str] = {
    "NOMBRE_EMPLEADO", "IDENTIDAD_EMPLEADO", "FECHA", "MARCA", "MODELO",
    "COSTO", "NUMERO_LINEA", "IMEI", "IDENTIFICADOR", "NOMBRE_USUARIO",
    "PUESTO", "IDENTIDAD_USUARIO", "OBSERVACION1", "OBSERVACION2", "MESES",
    "FIRMA_USUARIO", "FIRMA_EMPLEADO",
}

PC_FIELDS: Set[str] = {
    "FECHA", "NOMBRE_USUARIO", "IDENTIDAD_USUARIO", "IDENTIDAD_EMPLEADO",
    "PUESTO", "NOMBRE_EMPLEADO", "MARCA", "MODELO", "NUMERO_SERIE",
    "IDENTIFICADOR", "PROCESADOR", "RAM", "ALMACENAMIENTO", "OS",
    "TAMANO", "CARGADOR", "FIRMA_USUARIO", "FIRMA_EMPLEADO",
}

PERIFERICOS_FIELDS: Set[str] = {
    "FECHA", "NOMBRE_EMPLEADO", "IDENTIDAD_EMPLEADO", "NOMBRE_USUARIO",
    "IDENTIDAD_USUARIO", "PUESTO", "CATEGORIA", "MARCA", "MODELO",
    "NUMERO_SERIE", "IDENTIFICADOR", "FIRMA_USUARIO", "FIRMA_EMPLEADO",
}

CELULAR_AUTO_FIELDS: Set[str] = {"MARCA", "MODELO", "COSTO", "NUMERO_LINEA", "IMEI"}
LAPTOP_AUTO_FIELDS: Set[str] = {"MARCA", "MODELO", "NUMERO_SERIE", "PROCESADOR", "RAM", "ALMACENAMIENTO", "TAMANO", "CARGADOR", "OS"}
TABLET_AUTO_FIELDS: Set[str] = CELULAR_AUTO_FIELDS

CELULAR_MANUAL_FIELDS: Set[str] = TELEFONO_FIELDS - CELULAR_AUTO_FIELDS
LAPTOP_MANUAL_FIELDS: Set[str] = PC_FIELDS - LAPTOP_AUTO_FIELDS
TABLET_MANUAL_FIELDS: Set[str] = TELEFONO_FIELDS - TABLET_AUTO_FIELDS

_SIGNATURE_FIELDS: Set[str] = {"FIRMA_USUARIO", "FIRMA_EMPLEADO"}


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------
@dataclass
class DocFileBytes:
    name: str
    content: bytes


@dataclass
class DocExportResult:
    files: List[DocFileBytes]
    missing_fields: Tuple[str, ...]


# ---------------------------------------------------------------------------
# Paragraph iteration helpers  (eliminan el patrón 4-nivel que se repetía 6+ veces)
# ---------------------------------------------------------------------------

def _paragraphs_in_container(container) -> Iterator:
    """Itera párrafos de un contenedor (body, header, footer) incluyendo tablas."""
    yield from container.paragraphs
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _iter_all_paragraphs(doc) -> Iterator:
    """Itera TODOS los párrafos del documento: body + tablas + headers + footers."""
    yield from _paragraphs_in_container(doc)
    for section in doc.sections:
        yield from _paragraphs_in_container(section.header)
        yield from _paragraphs_in_container(section.footer)


# ---------------------------------------------------------------------------
# Core text / image replacement
# ---------------------------------------------------------------------------

def _find_run_span(runs_text: List[str], start_index: int, end_index: int):
    """Devuelve (start_run, start_off, end_run, end_off) para un rango de caracteres."""
    cum = 0
    start_run = start_off = 0
    for i, txt in enumerate(runs_text):
        if cum + len(txt) > start_index:
            start_run, start_off = i, start_index - cum
            break
        cum += len(txt)

    cum2 = cum
    end_run = start_run
    end_off = 0
    for j in range(start_run, len(runs_text)):
        if cum2 + len(runs_text[j]) >= end_index:
            end_run, end_off = j, end_index - cum2
            break
        cum2 += len(runs_text[j])

    return start_run, start_off, end_run, end_off


def _replace_in_paragraph(p, placeholder: str, value: str) -> None:
    """Reemplaza *placeholder* en un párrafo preservando formato de los runs.
    
    Cuando el placeholder se divide entre múltiples runs:
      - Coloca valor en primer run (preserva su formato)
      - Vacía runs intermedios sin eliminarlos (preserva estructura)
      - Mantiene suffix en último run
    """
    runs = list(p.runs)
    if not runs:
        return
    
    runs_text = [r.text or "" for r in runs]
    full_text = "".join(runs_text)
    if placeholder not in full_text:
        return

    idx = full_text.find(placeholder)
    while idx != -1:
        sr, so, er, eo = _find_run_span(runs_text, idx, idx + len(placeholder))
        prefix = runs_text[sr][:so]
        suffix = runs_text[er][eo:]

        if sr == er:
            # Placeholder en un solo run: preserva 100% del formato
            runs[sr].text = prefix + value + suffix
        else:
            # Placeholder dividido entre múltiples runs
            # Estrategia: preservar formato del primer run, vaciar intermedios
            runs[sr].text = prefix + value
            
            # Vaciar runs intermedios SIN eliminarlos (preserva estructura XML)
            for k in range(sr + 1, er):
                if runs[k].text:
                    runs[k].text = ""
            
            # Mantener suffix en último run con su formato
            runs[er].text = suffix

        # Refrescar para buscar la siguiente ocurrencia
        runs = list(p.runs)
        runs_text = [r.text or "" for r in runs]
        full_text = "".join(runs_text)
        idx = full_text.find(placeholder, idx + 1)


def _insert_image_in_paragraph_at(paragraph, run_index: int, image_stream, width_mm=None) -> None:
    new_run = paragraph.add_run()
    try:
        paragraph.runs[run_index]._element.addnext(new_run._element)
    except (IndexError, AttributeError) as e:
        logger.debug(f"No se pudo insertar imagen en índice {run_index}: {e}")

    if isinstance(image_stream, (bytes, bytearray)):
        # Usar BytesIO directamente en lugar de crear archivo temporal
        image_io = BytesIO(image_stream)
        try:
            new_run.add_picture(image_io, **({} if width_mm is None else {"width": Mm(width_mm)}))
        except Exception as e:
            logger.error(f"Error insertando imagen: {e}")
            new_run.text = "[IMAGEN]"
    else:
        try:
            new_run.add_picture(image_stream, **({} if width_mm is None else {"width": Mm(width_mm)}))
        except Exception as e:
            logger.error(f"Error insertando imagen: {e}")
            new_run.text = "[IMAGEN]"


def replace_image_placeholders(doc, images_map: Dict[str, Any], width_mm: int = 40) -> None:
    if not images_map:
        return
    placeholders = list(images_map.keys())

    def _process(p):
        runs = list(p.runs)
        if not runs:
            return
        runs_text = [r.text or "" for r in runs]
        full_text = "".join(runs_text)

        for key in placeholders:
            placeholder = f"[{key}]"
            idx = full_text.find(placeholder)
            while idx != -1:
                end_index = idx + len(placeholder)
                sr, so, er, eo = _find_run_span(runs_text, idx, end_index)
                prefix, suffix = runs_text[sr][:so], runs_text[er][eo:]

                runs[sr].text = prefix
                for k in range(sr + 1, er + 1):
                    runs[k].text = ""

                _insert_image_in_paragraph_at(p, sr, images_map[key], width_mm=width_mm)

                if suffix:
                    runs = list(p.runs)
                    if len(runs) > sr + 1:
                        runs[sr + 1].text = suffix
                    else:
                        p.add_run().text = suffix

                runs = list(p.runs)
                runs_text = [r.text or "" for r in runs]
                full_text = "".join(runs_text)
                idx = full_text.find(placeholder, idx + 1)

    for p in _iter_all_paragraphs(doc):
        _process(p)


# ---------------------------------------------------------------------------
# Public compat alias
# ---------------------------------------------------------------------------

def replace_placeholders(doc, replacements) -> None:
    """API pública: delega a docx_common.replace_placeholders."""
    logo_path = Path(__file__).parent.parent / "static" / "img" / "LogoProima.png"
    opts = PlaceholderOptions(logo_path=logo_path if logo_path.exists() else None)
    _replace_placeholders(doc, replacements, options=opts)


# ---------------------------------------------------------------------------
# Field helpers
# ---------------------------------------------------------------------------

def _ensure_fecha(fields: Dict[str, Any]) -> Dict[str, Any]:
    if fields.get("FECHA"):
        return fields
    return {**fields, "FECHA": datetime.now().strftime("%d/%m/%Y")}


def _build_token_map(fields: Dict[str, Any], required: Set[str]) -> Tuple[Dict[str, str], Tuple[str, ...]]:
    missing: List[str] = []
    token_map: Dict[str, str] = {}
    for k in sorted(required):
        val = fields.get(k)
        if val is None and k not in fields:
            missing.append(k)
        elif val is None:
            missing.append(k)
        else:
            token_map[k] = str(val).strip() if val != "" else ""
    if missing:
        logger.warning(f"_build_token_map: Missing fields: {missing}")
    return token_map, tuple(missing)


def _extraer_numero_correlativo(correlativo_completo: str) -> str:
    if not correlativo_completo:
        return "000000"
    try:
        partes = correlativo_completo.split("-")
        return partes[-1].zfill(6) if len(partes) >= 5 else "000000"
    except Exception:
        return "000000"


def _get_correlativo(fields: Dict[str, Any], prefix: str) -> str:
    """Obtiene el número de correlativo (6 dígitos) para un prefix dado."""
    correlativos = fields.get("CORRELATIVOS") or {}
    if not correlativos:
        fallback = fields.get("CORRELATIVO", "000000")
        correlativos = {prefix: f"{prefix}-{fallback}"}
        logger.warning(f"No se recibió CORRELATIVOS para {prefix}, usando fallback")
    completo = correlativos.get(prefix, f"{prefix}-000000")
    return _extraer_numero_correlativo(completo)


# ---------------------------------------------------------------------------
# DOCX serialization
# ---------------------------------------------------------------------------

def _save_tmp_docx(doc) -> bytes:
    """Guarda *doc* en un archivo temporal, valida el ZIP y devuelve bytes."""
    uid = str(uuid.uuid4())[:8]
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=f"_{uid}.docx", prefix="doc_")
    try:
        os.close(tmp_fd)
        doc.save(tmp_path)
        time.sleep(0.15)  # flush garantizado

        size = os.path.getsize(tmp_path)
        if size == 0:
            raise RuntimeError(f"Archivo temporal vacío: {tmp_path}")

        with zipfile.ZipFile(tmp_path, "r") as zf:
            for req in ("[Content_Types].xml", "word/document.xml"):
                if req not in zf.namelist():
                    raise RuntimeError(f"DOCX inválido: falta {req}")

        content = Path(tmp_path).read_bytes()
        if not content:
            raise RuntimeError(f"No se pudieron leer bytes de {tmp_path}")
        logger.debug(f"Documento generado: {len(content)} bytes")
        return content
    finally:
        try:
            os.unlink(tmp_path)
        except OSError as e:
            logger.warning(f"No se pudo eliminar tmp {tmp_path}: {e}")


def _apply_token_map(doc, token_map: Dict[str, str]) -> None:
    """Aplica todos los reemplazos de texto del token_map al documento."""
    for field_name, value in token_map.items():
        placeholder = f"[{field_name}]"
        for p in _iter_all_paragraphs(doc):
            _replace_in_paragraph(p, placeholder, value)


def _clear_remaining_placeholders(doc, required_fields: Set[str], token_map: Dict[str, str]) -> None:
    """Elimina placeholders de campos que no están en token_map."""
    for field_name in required_fields:
        if field_name not in token_map:
            placeholder = f"[{field_name}]"
            for p in _iter_all_paragraphs(doc):
                if placeholder in p.text:
                    _replace_in_paragraph(p, placeholder, "")


def _render_docx_bytes(
    template_path: Path,
    fields: Dict[str, Any],
    required_fields: Set[str],
    images_map: Optional[Dict[str, Any]] = None,
    image_width_mm: int = 40,
) -> Tuple[bytes, Tuple[str, ...]]:
    fields = _ensure_fecha(fields)
    token_map, missing = _build_token_map(fields, required_fields)

    # Inyectar CORRELATIVO si viene en fields
    corr = fields.get("CORRELATIVO")
    if "CORRELATIVO" not in token_map and corr is not None:
        token_map["CORRELATIVO"] = str(corr).zfill(6) if str(corr).isdigit() else str(corr)

    doc = Document(str(template_path))
    _apply_token_map(doc, token_map)

    if images_map:
        replace_image_placeholders(doc, images_map, width_mm=image_width_mm)

    _clear_remaining_placeholders(doc, required_fields, token_map)

    return _save_tmp_docx(doc), missing


def _render_docx_bytes_without_signatures(
    template_path: Path,
    fields: Dict[str, Any],
    required_fields: Set[str],
    images_map: Optional[Dict[str, Any]] = None,
    image_width_mm: int = 40,
) -> Tuple[bytes, Tuple[str, ...]]:
    """Primera pasada: reemplaza todo excepto los campos de firma."""
    fields = _ensure_fecha(fields)
    token_map, missing = _build_token_map(fields, required_fields)

    corr = fields.get("CORRELATIVO")
    if "CORRELATIVO" not in token_map and corr is not None:
        token_map["CORRELATIVO"] = str(corr).zfill(6) if str(corr).isdigit() else str(corr)

    doc = Document(str(template_path))

    # Si no hay images_map (firma manual), limpiar placeholders de firma
    if images_map is None:
        for sig in _SIGNATURE_FIELDS:
            token_map.setdefault(sig, "")

    for field_name, value in token_map.items():
        if field_name in _SIGNATURE_FIELDS:
            if images_map is not None:
                continue  # se insertarán en la segunda pasada
            value = ""
        placeholder = f"[{field_name}]"
        for p in _iter_all_paragraphs(doc):
            _replace_in_paragraph(p, placeholder, value)

    if images_map:
        non_sig = {k: v for k, v in images_map.items() if k not in _SIGNATURE_FIELDS}
        if non_sig:
            replace_image_placeholders(doc, non_sig, width_mm=image_width_mm)

    return _save_tmp_docx(doc), tuple(missing)


def _render_docx_bytes_signatures_only(docx_bytes_input: bytes, images_map: Dict[str, Any]) -> bytes:
    """Segunda pasada: inserta sólo las firmas en un DOCX ya generado."""
    doc = Document(BytesIO(docx_bytes_input))

    for field_name in _SIGNATURE_FIELDS:
        image_data = images_map.get(field_name)
        if not image_data:
            continue
        placeholder = f"[{field_name}]"

        for p in _iter_all_paragraphs(doc):
            if placeholder not in p.text:
                continue
            _replace_in_paragraph(p, placeholder, "")
            try:
                if isinstance(image_data, dict) and "path" in image_data:
                    p.add_run().add_picture(image_data["path"], width=Inches(2))
                elif isinstance(image_data, bytes):
                    p.add_run().add_picture(BytesIO(image_data), width=Inches(2))
            except Exception:
                logger.warning(f"No se pudo insertar imagen para {field_name}")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ---------------------------------------------------------------------------
# Export helpers
# ---------------------------------------------------------------------------

def _copy_images_map(images_map: Optional[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    if not images_map:
        return None
    return {
        k: bytes(v) if isinstance(v, (bytes, bytearray)) else v
        for k, v in images_map.items()
    }


def _export_single(
    prefix: str,
    template: Path,
    title: str,
    req_fields: Set[str],
    fields: Dict[str, Any],
    images_map: Optional[Dict[str, Any]],
) -> DocExportResult:
    """Exporta un único documento."""
    correlativo_numero = _get_correlativo(fields, prefix)
    fields_with_corr = {**fields, "CORRELATIVO": correlativo_numero}
    data, missing = _render_docx_bytes(template, fields_with_corr, req_fields, images_map=images_map)
    filename = f"{prefix}-{correlativo_numero} {title}.docx"
    return DocExportResult(files=[DocFileBytes(filename, data)], missing_fields=missing)


def export_celular(fields: Dict[str, Any], images_map: Optional[Dict[str, Any]] = None) -> DocExportResult:
    """Genera 2 documentos para Celular."""
    docs = [
        ("PRO-TI-CE-001", TEMPLATE_PRO_TI_001, "CERTIFICADO DE COMPROMISO Y ENTREGA DE TELEFONO CORPORATIVO", TELEFONO_FIELDS),
        ("PRO-TI-CE-002", TEMPLATE_PRO_TI_002, "MEMORANDO DE ENTREGA", TELEFONO_FIELDS),
    ]
    files: List[DocFileBytes] = []
    all_missing: List[str] = []

    for idx, (prefix, tpl, title, req) in enumerate(docs, 1):
        logger.info(f"=== DOCUMENTO {idx}/{len(docs)}: {prefix} ===")
        img_copy = _copy_images_map(images_map)
        gc.collect()

        correlativo_numero = _get_correlativo(fields, prefix)
        fields_with_corr = {**fields, "CORRELATIVO": correlativo_numero}
        data, missing = _render_docx_bytes(tpl, fields_with_corr, req, images_map=img_copy)

        if not data:
            raise RuntimeError(f"Documento {prefix} generó bytes vacíos")
        logger.info(f"✓ {prefix}: {len(data)} bytes")

        files.append(DocFileBytes(f"{prefix}-{correlativo_numero} {title}.docx", data))
        all_missing.extend(missing)

        gc.collect()
        time.sleep(0.3)

    return DocExportResult(files=files, missing_fields=tuple(sorted(set(all_missing))))


def export_laptop(fields: Dict[str, Any], images_map: Optional[Dict[str, Any]] = None) -> DocExportResult:
    return _export_single("PRO-TI-CE-004", TEMPLATE_PRO_TI_004, "CERTIFICADO ENTREGA DE COMPUTADORA", PC_FIELDS, fields, images_map)


def export_tablet(fields: Dict[str, Any], images_map: Optional[Dict[str, Any]] = None) -> DocExportResult:
    return _export_single("PRO-TI-CE-003", TEMPLATE_PRO_TI_003, "ENTREGA DE TABLET", TELEFONO_FIELDS, fields, images_map)


def export_periferico(fields: Dict[str, Any], images_map: Optional[Dict[str, Any]] = None) -> DocExportResult:
    return _export_single("PRO-TI-CE-005", TEMPLATE_PRO_TI_005, "ENTREGA DE PERIFERICO", PERIFERICOS_FIELDS, fields, images_map)


def docexport(categoria: str, fields: Dict[str, Any], images_map: Optional[Dict[str, Any]] = None) -> DocExportResult:
    tipo = (categoria or "").strip()
    if tipo in TIPOS_TELEFONO:
        return export_celular(fields, images_map=images_map)
    if tipo in TIPOS_PC:
        return export_laptop(fields, images_map=images_map)
    if tipo in TIPOS_TABLET:
        return export_tablet(fields, images_map=images_map)
    if tipo in TIPOS_PERIFERICOS:
        return export_periferico(fields, images_map=images_map)
    raise ValueError(f"Tipo no soportado: {categoria}")


def select_template(categoria: str) -> Path:
    tipo = (categoria or "").strip()
    mapping = {
        **{t: TEMPLATE_PRO_TI_001 for t in TIPOS_TELEFONO},
        **{t: TEMPLATE_PRO_TI_004 for t in TIPOS_PC},
        **{t: TEMPLATE_PRO_TI_003 for t in TIPOS_TABLET},
        **{t: TEMPLATE_PRO_TI_005 for t in TIPOS_PERIFERICOS},
    }
    if tipo not in mapping:
        raise ValueError(f"Tipo no soportado: {categoria}")
    return mapping[tipo]


# ---------------------------------------------------------------------------
# PDF conversion
# ---------------------------------------------------------------------------

def convert_docx_to_pdf(input_docx: str, output_pdf: Optional[str] = None) -> Optional[str]:
    input_path = Path(input_docx)
    if not input_path.exists():
        logger.error(f"convert_docx_to_pdf: Archivo no existe: {input_path}")
        return None
    if not HAS_DOCX2PDF:
        logger.warning("convert_docx_to_pdf: docx2pdf no está instalado")
        return None

    out_path = Path(output_pdf) if output_pdf else input_path.with_suffix(".pdf")
    try:
        time.sleep(0.2)
        _docx2pdf_convert(str(input_path), str(out_path))
        if out_path.exists():
            logger.info(f"PDF creado: {out_path.name} ({out_path.stat().st_size} bytes)")
            return str(out_path)
        logger.error(f"No se creó el PDF: {out_path}")
    except Exception as e:
        logger.error(f"Error convirtiendo {input_path.name}: {e}", exc_info=True)
    return None


def _save_bytes(path: Path, data: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(data)


def save_export_result(result: DocExportResult, output_dir: str | Path) -> Tuple[Tuple[str, ...], Tuple[str, ...]]:
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    paths = []
    for f in result.files:
        target = out_dir / f.name
        _save_bytes(target, f.content)
        paths.append(str(target))
    return tuple(paths), result.missing_fields


def convert_many_docx_to_pdf(docx_paths: List[str], output_dir: str | Path) -> Tuple[Tuple[str, ...], Tuple[str, ...]]:
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    pdfs, failed = [], []
    for idx, docx in enumerate(docx_paths, 1):
        docx_p = Path(docx)
        pdf_target = out_dir / (docx_p.stem + ".pdf")
        logger.info(f"[{idx}/{len(docx_paths)}] {docx_p.name} -> {pdf_target.name}")
        pdf_path = convert_docx_to_pdf(str(docx_p), str(pdf_target))
        (pdfs if pdf_path else failed).append(pdf_path or str(docx_p))
    logger.info(f"PDF: {len(pdfs)} ok, {len(failed)} fallidos")
    return tuple(pdfs), tuple(failed)


def export_and_convert(
    categoria: str,
    fields: Dict[str, Any],
    output_docx_dir: str | Path,
    output_pdf: Optional[str] = None,
    images_map: Optional[Dict[str, Any]] = None,
) -> Tuple[DocExportResult, Optional[str]]:
    result = docexport(categoria, fields, images_map=images_map)
    docx_paths, _ = save_export_result(result, output_docx_dir)
    if not docx_paths:
        return result, None
    first_docx = docx_paths[0]
    pdf_target = Path(output_pdf) if output_pdf else Path(output_docx_dir) / (Path(first_docx).stem + ".pdf")
    return result, convert_docx_to_pdf(first_docx, str(pdf_target))


def export_and_convert_many(
    categoria: str,
    fields: Dict[str, Any],
    output_docx_dir: str | Path,
    output_pdf_dir: Optional[str | Path] = None,
    images_map: Optional[Dict[str, Any]] = None,
) -> Tuple[DocExportResult, Tuple[str, ...], Tuple[str, ...]]:
    result = docexport(categoria, fields, images_map=images_map)
    docx_paths, _ = save_export_result(result, output_docx_dir)
    pdf_dir = output_pdf_dir if output_pdf_dir is not None else output_docx_dir
    pdf_paths, failed = convert_many_docx_to_pdf(list(docx_paths), pdf_dir)
    return result, pdf_paths, failed


def zip_export_result(result: DocExportResult, zip_name: str = "documentos.zip") -> bytes:
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for f in result.files:
            zf.writestr(f.name, f.content)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# REST API Blueprint
# ---------------------------------------------------------------------------
from flask import Blueprint, request, jsonify

docexp_bp = Blueprint("docexp", __name__, url_prefix="/api/docexp")

# Configuración de cada categoría para los endpoints generados dinámicamente
_CATEGORIA_CONFIG = {
    "celular": {
        "categoria": "Celular",
        "export_fn": export_celular,
        "manual_fields": CELULAR_MANUAL_FIELDS,
        "auto_fields": CELULAR_AUTO_FIELDS,
        "all_fields": TELEFONO_FIELDS,
        "templates": [
            "PRO-TI-CE-001 (Certificado de Compromiso Teléfono)",
            "PRO-TI-CE-002 (Memorando de Entrega)",
        ],
        "field_sources": {
            "MARCA": "dispositivo.nombre_marca",
            "MODELO": "dispositivo.nombre_modelo",
            "COSTO": "asignacion.costo_plan",
            "NUMERO_LINEA": "asignacion.numero_linea (desde plan)",
            "IMEI": "dispositivo.imei",
        },
    },
    "laptop": {
        "categoria": "Laptop",
        "export_fn": export_laptop,
        "manual_fields": LAPTOP_MANUAL_FIELDS,
        "auto_fields": LAPTOP_AUTO_FIELDS,
        "all_fields": PC_FIELDS,
        "templates": ["PRO-TI-CE-004 (Certificado Entrega de Computadora)"],
        "field_sources": {
            "MARCA": "dispositivo.nombre_marca",
            "MODELO": "dispositivo.nombre_modelo",
            "NUMERO_SERIE": "dispositivo.numero_serie",
            "PROCESADOR": "componente CPU (marca + modelo)",
            "RAM": "componente RAM (capacidad)",
            "ALMACENAMIENTO": "componente DISCO (capacidad)",
            "TAMANO": "dispositivo.tamano",
            "CARGADOR": "dispositivo.cargador",
            "OS": "componente CPU.observaciones (Sistema Operativo)",
        },
    },
    "tablet": {
        "categoria": "Tablet",
        "export_fn": export_tablet,
        "manual_fields": TABLET_MANUAL_FIELDS,
        "auto_fields": TABLET_AUTO_FIELDS,
        "all_fields": TELEFONO_FIELDS,
        "templates": ["PRO-TI-CE-006 (Entrega de Tablet)"],
        "field_sources": {
            "MARCA": "dispositivo.nombre_marca",
            "MODELO": "dispositivo.nombre_modelo",
            "COSTO": "asignacion.costo_plan",
            "NUMERO_LINEA": "asignacion.numero_linea (desde plan)",
            "IMEI": "dispositivo.imei",
        },
    },
    "periferico": {
        "categoria": "Periférico",
        "export_fn": export_periferico,
        "required_fields": PERIFERICOS_FIELDS,
        "tipos_soportados": TIPOS_PERIFERICOS,
        "templates": ["PRO-TI-CE-005 (Entrega de Periférico)"],
    },
}


def _result_to_json_files(result: DocExportResult) -> List[Dict]:
    return [
        {"name": f.name, "content": base64.b64encode(f.content).decode("utf-8"), "size": len(f.content)}
        for f in result.files
    ]


def _generate_endpoint(slug: str):
    cfg = _CATEGORIA_CONFIG[slug]
    data = request.get_json(silent=True) or {}
    fields = data.get("fields", {})
    if not fields:
        return jsonify({"success": False, "message": 'El campo "fields" es requerido'}), 400
    try:
        result = cfg["export_fn"](fields, images_map=data.get("images_map"))
        return jsonify({
            "success": True,
            "categoria": cfg["categoria"],
            "files": _result_to_json_files(result),
            "missing_fields": list(result.missing_fields),
            "files_count": len(result.files),
        })
    except Exception as e:
        logger.exception(f"Error generando documentos para {cfg['categoria']}")
        return jsonify({"success": False, "message": f"Error: {e}"}), 500


def _fields_endpoint(slug: str):
    cfg = _CATEGORIA_CONFIG[slug]
    payload: Dict[str, Any] = {"success": True, "categoria": cfg["categoria"]}
    if "manual_fields" in cfg:
        payload["manual_fields"] = sorted(cfg["manual_fields"])
        payload["auto_fields"] = sorted(cfg["auto_fields"])
        payload["all_fields"] = sorted(cfg["all_fields"])
    if "required_fields" in cfg:
        payload["required_fields"] = sorted(cfg["required_fields"])
    if "tipos_soportados" in cfg:
        payload["tipos_soportados"] = sorted(cfg["tipos_soportados"])
    payload["templates"] = cfg["templates"]
    if "field_sources" in cfg:
        payload["field_sources"] = cfg["field_sources"]
    return jsonify(payload)


# Registro de rutas
for _slug in _CATEGORIA_CONFIG:
    # Usar closures para capturar _slug correctamente
    def _make_fields_view(s):
        def view():
            return _fields_endpoint(s)
        view.__name__ = f"get_{s}_fields"
        return view

    def _make_generate_view(s):
        def view():
            return _generate_endpoint(s)
        view.__name__ = f"generate_{s}_docs"
        return view

    docexp_bp.add_url_rule(f"/{_slug}/fields", view_func=_make_fields_view(_slug), methods=["GET"])
    docexp_bp.add_url_rule(f"/{_slug}/generate", view_func=_make_generate_view(_slug), methods=["POST"])


@docexp_bp.get("/health")
def health_check():
    return jsonify({
        "success": True,
        "service": "docexp",
        "status": "healthy",
        "has_docx2pdf": HAS_DOCX2PDF,
        "categorias_soportadas": [cfg["categoria"] for cfg in _CATEGORIA_CONFIG.values()],
    })
