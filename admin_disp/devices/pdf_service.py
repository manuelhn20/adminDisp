"""
DEPRECATED: Este módulo ha sido consolidado en admin_disp/services/docgen.DocumentGenerator.

Para generar documentos, usa:
    from admin_disp.services.docgen import DocumentGenerator
    doc_gen = DocumentGenerator()
    pdf_buffer = doc_gen.generate_for_assignment(asignacion_id, coordinador_data, empleado_data)

PDFGeneratorService se mantiene aquí solo por compatibilidad backward. No usar en nuevo código.
"""

import logging
import os
import json
import tempfile
import shutil
import subprocess
from datetime import datetime
from io import BytesIO
from pathlib import Path
import re

logger = logging.getLogger(__name__)

from admin_disp.services.docform import replace_placeholders
from admin_disp.services.docx_common import RemapData


class PDFGeneratorService:
    """Genera PDFs de entrega de dispositivos copiando y modificando el documento Word original."""

    @staticmethod
    def _resolve_template_word_path() -> str:
        base_dir = Path(__file__).resolve().parent.parent / "form"
        candidates = [
            "PRO-TI-CE-001-CORRELATIVO CERTIFICADO DE COMPROMISO Y ENTREGA DE TELEFONO CORPORATIVO.docx",
            "PRO-TI-CE-003-CORRELATIVO CERTIFICADO DE COMPROMISO Y ENTREGA DE TABLET.docx",
            "PRO-TI-CE-003-###### ENTREGA DE TELEFONO.docx",
        ]
        for name in candidates:
            candidate = base_dir / name
            if candidate.exists():
                return str(candidate)
        return str(base_dir / candidates[0])

    # Ruta al documento Word template (compatibilidad con nombres antiguos y nuevos)
    TEMPLATE_WORD_PATH = _resolve_template_word_path.__func__()
    
    def __init__(self):
        self._check_libreoffice()

    def _build_pc_replacements(self, dispositivo: dict, componentes: list) -> dict:
        """Construye replacements específicos para plantillas PC/Laptop.

        Retorna un dict con claves: PROCESADOR, RAM, ALMACENAMIENTO, OS, TAMANO, CARGADOR
        """
        out = {}
        try:
            # PROCESADOR: intentar extraer marca+modelo del componente CPU
            proc = None
            if componentes:
                for c in componentes:
                    try:
                        if str(c.get('tipo_componente')).upper() == 'CPU':
                            proc = c
                            break
                    except Exception:
                        pass

            if proc:
                marca_cpu = proc.get('nombre_marca') or proc.get('marca') or dispositivo.get('nombre_marca') or dispositivo.get('MARCA') or ''
                # Only read modelo from the componente itself (do NOT fallback to dispositivo yet)
                modelo_cpu = proc.get('modelo') or proc.get('nombre_modelo') or ''

                # If the proc component does not contain a useful modelo, first check if the
                # componente itself references a modelo via `fk_id_modelo` and use it.
                if not modelo_cpu:
                    try:
                        if proc.get('fk_id_modelo'):
                            from admin_disp.devices.service import DeviceService
                            svc_tmp = DeviceService()
                            mtmp = svc_tmp.get_modelo(int(proc.get('fk_id_modelo')))
                            if mtmp and mtmp.get('nombre_modelo'):
                                modelo_cpu = mtmp.get('nombre_modelo')
                    except Exception:
                        pass

                # If still empty, prefer a modelo stored in other componentes (some modelos
                # are stored as componentes with estado==2).
                if not modelo_cpu and componentes:
                    try:
                        # First prefer componentes with fk_id_modelo and estado==2
                        from admin_disp.devices.service import DeviceService
                        svc = DeviceService()
                        found = False
                        for c in componentes:
                            try:
                                if c.get('estado') in (2, '2') and c.get('fk_id_modelo'):
                                    m = svc.get_modelo(int(c.get('fk_id_modelo')))
                                    if m and m.get('nombre_modelo'):
                                        modelo_cpu = m.get('nombre_modelo')
                                        found = True
                                        break
                            except Exception:
                                pass
                        # If none with estado==2, try any componente with fk_id_modelo
                        if not found:
                            for c in componentes:
                                try:
                                    if c.get('fk_id_modelo'):
                                        m = svc.get_modelo(int(c.get('fk_id_modelo')))
                                        if m and m.get('nombre_modelo'):
                                            modelo_cpu = m.get('nombre_modelo')
                                            break
                                except Exception:
                                    pass
                    except Exception:
                        pass

                # If still empty after checking componente and other componentes, fallback to dispositivo model
                if not modelo_cpu and componentes:
                    try:
                        # First prefer componentes with fk_id_modelo and estado==2
                        from admin_disp.devices.service import DeviceService
                        svc = DeviceService()
                        found = False
                        for c in componentes:
                            try:
                                if c.get('estado') in (2, '2') and c.get('fk_id_modelo'):
                                    m = svc.get_modelo(int(c.get('fk_id_modelo')))
                                    if m and m.get('nombre_modelo'):
                                        modelo_cpu = m.get('nombre_modelo')
                                        found = True
                                        break
                            except Exception:
                                pass
                        # If none with estado==2, try any componente with fk_id_modelo
                        if not found:
                            for c in componentes:
                                try:
                                    if c.get('fk_id_modelo'):
                                        m = svc.get_modelo(int(c.get('fk_id_modelo')))
                                        if m and m.get('nombre_modelo'):
                                            modelo_cpu = m.get('nombre_modelo')
                                            break
                                except Exception:
                                    pass
                    except Exception:
                        pass

                proc_str = (marca_cpu + ' ' + (modelo_cpu or '')).strip()
                if not proc_str:
                    proc_str = (dispositivo.get('nombre_marca') or '') + ' ' + (dispositivo.get('nombre_modelo') or '')
                out['PROCESADOR'] = proc_str.strip()
            else:
                out['PROCESADOR'] = ((dispositivo.get('nombre_marca') or '') + ' ' + (dispositivo.get('nombre_modelo') or '')).strip()

            # RAM: sumar capacidades de módulos RAM si existen, si no usar campo dispositivo si existe
            ram_val = None
            if componentes:
                total_ram = 0
                any_ram = False
                for c in componentes:
                    try:
                        if str(c.get('tipo_componente')).upper() == 'RAM':
                            cap = c.get('capacidad')
                            if cap:
                                total_ram += int(cap)
                                any_ram = True
                    except Exception:
                        pass
                if any_ram:
                    ram_val = f"{total_ram} GB"

            if not ram_val:
                # fallback: intentar campos comunes en dispositivo
                dv_ram = dispositivo.get('ram') or dispositivo.get('RAM')
                if dv_ram:
                    try:
                        ram_val = f"{int(dv_ram)} GB"
                    except Exception:
                        ram_val = str(dv_ram)

            out['RAM'] = ram_val or ''

            # ALMACENAMIENTO: elegir disco principal (mayor capacidad) y mostrar solo capacidad
            disco_val = ''
            if componentes:
                best = None
                for c in componentes:
                    try:
                        if str(c.get('tipo_componente')).upper() in ('DISCO', 'DISK'):
                            cap = c.get('capacidad')
                            if cap is None:
                                continue
                            if best is None or int(cap) > int(best.get('capacidad') or 0):
                                best = c
                    except Exception:
                        pass
                if best:
                    disco_val = f"{int(best.get('capacidad'))} GB"

            if not disco_val:
                dv_storage = dispositivo.get('almacenamiento') or dispositivo.get('ALMACENAMIENTO')
                if dv_storage:
                    disco_val = str(dv_storage)

            out['ALMACENAMIENTO'] = disco_val or ''

            # OS
            out['OS'] = (dispositivo.get('os') or dispositivo.get('OS') or '')

            # TAMANO
            out['TAMANO'] = str(dispositivo.get('tamano') or dispositivo.get('TAMANO') or '')

            # CARGADOR: mostrar SI/NO
            carg = dispositivo.get('cargador')
            if carg in (1, '1', True, 'true', 'True'):
                out['CARGADOR'] = 'SI'
            elif carg in (0, '0', False, 'false', 'False'):
                out['CARGADOR'] = 'NO'
            else:
                # fallback: truthy check
                out['CARGADOR'] = 'SI' if carg else 'NO'

        except Exception:
            # en caso de error devolver keys vacías para evitar romper flujo
            out.setdefault('PROCESADOR', '')
            out.setdefault('RAM', '')
            out.setdefault('ALMACENAMIENTO', '')
            out.setdefault('OS', '')
            out.setdefault('TAMANO', '')
            out.setdefault('CARGADOR', '')

        return out
    
    def _check_libreoffice(self):
        """Verifica si LibreOffice está disponible en el sistema."""
        try:
            # Intenta ejecutar soffice --version
            subprocess.run(
                ['soffice', '--version'],
                capture_output=True,
                timeout=5,
                check=False
            )
            self.has_libreoffice = True
            logger.info("LibreOffice detectado")
        except:
            self.has_libreoffice = False
    
    def generate_entrega_telefono_pdf(self, asignacion_data: dict, dispositivo_data: dict, 
                                       empleado_data: dict, coordinador_data: dict, observaciones: list = None,
                                       template_path: str | None = None) -> BytesIO:
        """
        Genera un PDF de entrega de teléfono basado en los datos de la asignación.
        
        Copia el documento Word original, reemplaza los campos dinámicos y convierte a PDF.
        
        Parámetros:
        - asignacion_data: dict con datos de la asignación (fecha_inicio_asignacion, numero_linea, costo_plan)
        - dispositivo_data: dict con datos del dispositivo (nombre_modelo, nombre_marca, imei)
        - empleado_data: dict con datos del empleado (nombre_completo)
        - coordinador_data: dict con datos del coordinador (nombre)
        
        Retorna: BytesIO con el contenido del PDF
        """
        
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx no está instalado")
            raise RuntimeError("python-docx es requerido para generar PDFs")
        
        try:
            # Determinar ruta del template a usar (se puede suministrar una plantilla personalizada)
            tpl = template_path if template_path else self.TEMPLATE_WORD_PATH
            if not os.path.exists(tpl):
                raise FileNotFoundError(f"Documento template no encontrado: {tpl}")

            # Crear copia temporal del documento
            temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            temp_docx_path = temp_docx.name
            temp_docx.close()

            # Copiar el documento original
            shutil.copy2(tpl, temp_docx_path)
            
            # Abrir el documento copiado
            doc = Document(temp_docx_path)
            
            # Extraer datos para reemplazar
            nombre_empleado = empleado_data.get('nombre_completo') or empleado_data.get('NombreCompleto') or ''
            marca = dispositivo_data.get('nombre_marca', '')
            modelo = dispositivo_data.get('nombre_modelo', '')
            # Extraer moneda y costo, concatenarlos
            moneda = asignacion_data.get('moneda_plan', '$')
            costo = asignacion_data.get('costo_plan', '')
            costo_str = f"{moneda}{costo}" if costo else ''
            fecha = self._format_date(asignacion_data.get('fecha_inicio_asignacion'))
            linea = asignacion_data.get('numero_linea', '')
            # Limpiar posibles llaves literales y formatear la línea a "+(ext) XXXX-XXXX"
            try:
                if isinstance(linea, str):
                    # eliminar llaves si llegaron como {{Linea}}
                    linea_clean = linea.replace('{', '').replace('}', '').strip()
                    # extraer sólo dígitos
                    digits = re.sub(r'\D', '', linea_clean or '')
                    formatted_linea = linea_clean
                    if digits:
                        # si hay al menos 8 dígitos, tomar los últimos 8 como número y el resto como extensión
                        if len(digits) >= 8:
                            ext = digits[:-8]
                            num = digits[-8:]
                            num_formatted = f"{num[:4]}-{num[4:]}"
                            if ext:
                                # limitar extensión a 1-3 dígitos si es posible
                                ext = ext[-3:]
                                formatted_linea = f"+({ext}) {num_formatted}"
                            else:
                                formatted_linea = num_formatted
                        else:
                            # si no hay 8 dígitos, conservar los dígitos agrupados seguidos
                            formatted_linea = digits
                    linea = formatted_linea
                else:
                    linea = str(linea or '')
            except Exception:
                linea = asignacion_data.get('numero_linea', '')
            imei = dispositivo_data.get('imei', '')
            # Número de serie / serie del dispositivo
            serie = dispositivo_data.get('serie') or dispositivo_data.get('numero_serie') or dispositivo_data.get('numero_de_serie') or ''
            # No forzar "Coordinador IT" como nombre; el nombre debe venir en coordinador_data
            coordinador = coordinador_data.get('nombre', '') if isinstance(coordinador_data, dict) else ''

            # Usuario/coordinador
            nombre_usuario = (coordinador_data.get('nombre') if isinstance(coordinador_data, dict) else coordinador) or ''
            nombre_usuario = nombre_usuario.strip() if nombre_usuario else ''
            
            # Construir replacements: NO incluir NOMBRE_USUARIO si está vacío (mantener placeholder)
            replacements = {
                'NOMBRE_EMPLEADO': nombre_empleado,
                'IDENTIDAD_EMPLEADO': empleado_data.get('numero_identidad', ''),
                'IDENTIDAD_USUARIO': coordinador_data.get('numero_identidad', '') if isinstance(coordinador_data, dict) else '',
                # PUESTO: preferir puesto del coordinador; si no existe usar el valor por defecto
                'PUESTO': (coordinador_data.get('puesto') if isinstance(coordinador_data, dict) and coordinador_data.get('puesto') else 'Coordinador de IT'),
                'FECHA': fecha,
                'Fecha': fecha,
                'fecha': fecha,
                'MARCA': marca,
                'MODELO': modelo,
                'NUMERO_LINEA': linea,
                # Variantes comunes del placeholder para línea
                'Linea': linea,
                'LINEA': linea,
                'Numero_Linea': linea,
                'NUMERO_LINEA': linea,
                'IMEI': imei,
                'NUMERO_SERIE': serie,
                'COORDINADOR': coordinador,
                'COSTO': costo_str,
            }

            # Añadir NOMBRE_USUARIO solo si existe; si no, dejamos el placeholder tal cual
            if nombre_usuario:
                replacements['NOMBRE_USUARIO'] = nombre_usuario

            # Añadir variantes y claves comunes que pueden aparecer en plantillas (p.ej. CATEGORIA)
            try:
                tipo_disp_val = dispositivo_data.get('categoria') or dispositivo_data.get('tipo') or ''
            except Exception:
                tipo_disp_val = ''
            # Normalizar valores a strings
            tipo_disp_val = str(tipo_disp_val or '')

            extra_keys = {
                'CATEGORIA': tipo_disp_val,
                'TIPO': tipo_disp_val,
                'TIPO_DISP': tipo_disp_val,
                'TIPO DISPOSITIVO': tipo_disp_val,
                'TIPO-DISPOSITIVO': tipo_disp_val,
            }

            # Añadir las claves extra al diccionario replacements (no sobreescribir si ya existen)
            for k, v in extra_keys.items():
                if k not in replacements:
                    replacements[k] = v

            # Si es una Laptop/PC y vienen componentes adjuntos, construir replacements PC específicos
            try:
                tipo_disp_val = tipo_disp_val.strip()
                # importar TIPOS_PC para coincidencia exacta (usando docexp, módulo más actualizado)
                from admin_disp.services.docexp import TIPOS_PC
                comps = dispositivo_data.get('componentes') if isinstance(dispositivo_data, dict) else None
                if tipo_disp_val in TIPOS_PC and comps is not None:
                    pc_repls = self._build_pc_replacements(dispositivo_data, comps)
                    # Agregar sin sobreescribir valores existentes
                    for k, v in pc_repls.items():
                        if k and (k not in replacements or not replacements.get(k)):
                            replacements[k] = v
            except Exception:
                pass

            # También asegurar variantes en mayúsculas/minúsculas para nombres clave ya presentes
            additions = {}
            for k, v in list(replacements.items()):
                try:
                    # No generar variantes si el valor está vacío (evita reemplazar el placeholder por cadena vacía)
                    if not v:
                        continue
                    if k.upper() not in replacements:
                        additions[k.upper()] = v
                    if k.lower() not in replacements:
                        additions[k.lower()] = v
                except Exception:
                    pass
            replacements.update(additions)

            # Agregar observaciones (1..5). Soportar mayúsculas/minúsculas y variantes
            try:
                for i in range(1, 6):
                    # Si la observación fue enviada, usar su texto; si no, usar dos espacios zero-width
                    # para que docexport la considere 'no vacía' y reemplace el placeholder por un espacio invisible.
                    if observaciones and len(observaciones) >= i and observaciones[i-1]:
                        val = observaciones[i-1]
                    else:
                        val = '\u200B\u200B'
                    # Variantes sin espacio y con guion bajo
                    replacements[f'OBSERVACION_{i}'] = val
                    replacements[f'Observacion_{i}'] = val
                    replacements[f'observacion_{i}'] = val
                    # Variantes con espacio
                    replacements[f'OBSERVACION {i}'] = val
                    replacements[f'Observacion {i}'] = val
                    replacements[f'observacion {i}'] = val
                    # Variantes con tilde (acentuadas)
                    replacements[f'OBSERVACIÓN_{i}'] = val
                    replacements[f'Observación_{i}'] = val
                    replacements[f'observación_{i}'] = val
                    replacements[f'OBSERVACIÓN {i}'] = val
                    replacements[f'Observación {i}'] = val
                    replacements[f'observación {i}'] = val
            except Exception:
                pass

            try:
                    # Log: plantilla usada y diccionario de reemplazos (para depuración)
                    try:
                        logger.info(f"PDF template path used: {tpl}")
                        logger.info(f"Replacements keys: {list(replacements.keys())}")
                    except Exception:
                        pass

                    # Log: mostrar fragmentos del documento que contienen placeholders
                    try:
                        sample_hits = []
                        for p in doc.paragraphs:
                            txt = ''.join(run.text for run in p.runs)
                            if '[' in txt or '{{' in txt or '}}' in txt:
                                sample_hits.append(txt.strip())
                                if len(sample_hits) >= 8:
                                    break
                        # also scan first table cells
                        if len(sample_hits) < 8:
                            for t in getattr(doc, 'tables', []):
                                for row in t.rows:
                                    for cell in row.cells:
                                        for p in cell.paragraphs:
                                            txt = ''.join(run.text for run in p.runs)
                                            if '[' in txt or '{{' in txt or '}}' in txt:
                                                sample_hits.append(txt.strip())
                                            if len(sample_hits) >= 8:
                                                break
                                    if len(sample_hits) >= 8:
                                        break
                                if len(sample_hits) >= 8:
                                    break
                        logger.info(f"Document sample placeholders: {sample_hits}")
                    except Exception:
                        pass

                    # Primero aplicar la lógica estándar de `docexport` (reemplaza [KEY])
                    replace_placeholders(doc, replacements)
                    # Además, asegurar reemplazo de variantes con llaves dobles {{KEY}} y otras variantes
                    def _replace_curly_variants(document, repls):
                        def _replace_in_paragraph_runs(paragraph):
                            if not paragraph.runs:
                                return
                            for run in paragraph.runs:
                                text = run.text
                                if not text:
                                    continue
                                for kk, vv in repls.items():
                                    try:
                                        text = text.replace(f'{{{{{kk}}}}}', vv)
                                        text = text.replace(f'{{{{ {kk} }}}}', vv)
                                        text = text.replace(f'{{{kk}}}', vv)
                                    except Exception:
                                        pass
                                run.text = text

                        for p in document.paragraphs:
                            _replace_in_paragraph_runs(p)
                        for t in getattr(document, 'tables', []):
                            for row in t.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        _replace_in_paragraph_runs(p)
                        # headers/footers
                        for section in getattr(document, 'sections', []):
                            try:
                                for p in section.header.paragraphs:
                                    _replace_in_paragraph_runs(p)
                                for t in section.header.tables:
                                    for row in t.rows:
                                        for cell in row.cells:
                                            for p in cell.paragraphs:
                                                _replace_in_paragraph_runs(p)
                            except Exception:
                                pass
                            try:
                                for p in section.footer.paragraphs:
                                    _replace_in_paragraph_runs(p)
                                for t in section.footer.tables:
                                    for row in t.rows:
                                        for cell in row.cells:
                                            for p in cell.paragraphs:
                                                _replace_in_paragraph_runs(p)
                            except Exception:
                                pass

                    _replace_curly_variants(doc, replacements)
                    # Además, asegurar reemplazo directo de placeholders con corchetes [KEY]
                    def _replace_square_variants(document, repls):
                        def _replace_in_paragraph_runs(paragraph):
                            if not paragraph.runs:
                                return
                            for run in paragraph.runs:
                                text = run.text
                                if not text:
                                    continue
                                for kk, vv in repls.items():
                                    try:
                                        text = text.replace(f'[{kk}]', vv)
                                        text = text.replace(f'[ {kk} ]', vv)
                                        text = text.replace(f'[{kk} ]', vv)
                                        text = text.replace(f'[ {kk}]', vv)
                                    except Exception:
                                        pass
                                run.text = text

                        for p in document.paragraphs:
                            _replace_in_paragraph_runs(p)
                        for t in getattr(document, 'tables', []):
                            for row in t.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        _replace_in_paragraph_runs(p)
                        for section in getattr(document, 'sections', []):
                            try:
                                for p in section.header.paragraphs:
                                    _replace_in_paragraph_runs(p)
                                for t in section.header.tables:
                                    for row in t.rows:
                                        for cell in row.cells:
                                            for p in cell.paragraphs:
                                                _replace_in_paragraph_runs(p)
                            except Exception:
                                pass
                            try:
                                for p in section.footer.paragraphs:
                                    _replace_in_paragraph_runs(p)
                                for t in section.footer.tables:
                                    for row in t.rows:
                                        for cell in row.cells:
                                            for p in cell.paragraphs:
                                                _replace_in_paragraph_runs(p)
                            except Exception:
                                pass

                    _replace_square_variants(doc, replacements)

                    # Pasada adicional robusta: reemplazo de [KEY] sobre el texto completo
                    def _replace_square_brackets_full_paragraphs(document, repls):
                        import re as _re
                        # Build normalized key map: normalize keys to UPPER and underscores
                        def _norm(k):
                            if k is None:
                                return ''
                            kk = str(k).strip()
                            kk = kk.replace('-', ' ').replace(' ', '_')
                            return kk.upper()

                        norm_map = {}
                        for k, v in repls.items():
                            try:
                                norm_map[_norm(k)] = str(v)
                            except Exception:
                                norm_map[_norm(k)] = ''

                        token_re = _re.compile(r"\[([^\]]+)\]")

                        def _process_paragraph(paragraph):
                            full_text = ''.join(run.text for run in paragraph.runs)
                            if not full_text or ('[' not in full_text):
                                return False
                            changed = False
                            def _repl(m):
                                key = m.group(1)
                                n = _norm(key)
                                if n in norm_map:
                                    nonlocal_changed[0] = True
                                    return norm_map[n]
                                # try also direct key
                                if key in repls:
                                    nonlocal_changed[0] = True
                                    return str(repls[key])
                                return m.group(0)

                            nonlocal_changed = [False]
                            new_text = token_re.sub(_repl, full_text)
                            if nonlocal_changed[0] and new_text != full_text:
                                # replace runs: clear and set single run to new_text
                                for run in paragraph.runs:
                                    run.text = ''
                                if paragraph.runs:
                                    paragraph.runs[0].text = new_text
                                else:
                                    # create a run if none
                                    try:
                                        paragraph.add_run(new_text)
                                    except Exception:
                                        pass
                                return True
                            return False

                        # paragraphs
                        for p in document.paragraphs:
                            try:
                                _process_paragraph(p)
                            except Exception:
                                pass

                        # tables
                        for t in getattr(document, 'tables', []):
                            for row in t.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        try:
                                            _process_paragraph(p)
                                        except Exception:
                                            pass

                        # headers and footers
                        for section in getattr(document, 'sections', []):
                            try:
                                for p in section.header.paragraphs:
                                    try:
                                        _process_paragraph(p)
                                    except Exception:
                                        pass
                            except Exception:
                                pass
                            try:
                                for p in section.footer.paragraphs:
                                    try:
                                        _process_paragraph(p)
                                    except Exception:
                                        pass
                            except Exception:
                                pass

                    # Nota: no ejecutar la pasada destructiva que reemplaza párrafos completos
                    # porque rompe formato (negritas/estilos). Confiamos en `replace_placeholders`
                    # y en los reemplazos por-run (`_replace_curly_variants`, `_replace_square_variants`).
                    # _replace_square_brackets_full_paragraphs(doc, replacements)
            except Exception as e:
                logger.warning(f"Error aplicando replace_placeholders desde docexport: {e}")
            
            # Guardar el documento modificado
            temp_docx_out = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            temp_docx_out_path = temp_docx_out.name
            temp_docx_out.close()
            
            doc.save(temp_docx_out_path)
            
            # Convertir a PDF si LibreOffice está disponible
            if self.has_libreoffice:
                pdf_path = self._convert_docx_to_pdf(temp_docx_out_path)
                # Leer el PDF en BytesIO
                with open(pdf_path, 'rb') as pdf_file:
                    pdf_buffer = BytesIO(pdf_file.read())
            else:
                # Intentar usar docx2pdf vía el helper en `admin_disp.services.docexport`
                # si está instalado y disponible. Esto permite generar PDF en
                # entornos Windows donde Word está presente.
                pdf_path = None
                try:
                    from admin_disp.services import docexp as de
                    try:
                        pdf_path = de.convert_docx_to_pdf(temp_docx_out_path)
                    except Exception:
                        pdf_path = None
                except Exception:
                    pdf_path = None

                if pdf_path and os.path.exists(pdf_path):
                    try:
                        with open(pdf_path, 'rb') as pdf_file:
                            pdf_buffer = BytesIO(pdf_file.read())
                    except Exception:
                        # Fallback to returning DOCX bytes if reading PDF fails
                        with open(temp_docx_out_path, 'rb') as docx_file:
                            pdf_buffer = BytesIO(docx_file.read())
                else:
                    # Si no hay conversión disponible, devolver el DOCX modificado como BytesIO
                    with open(temp_docx_out_path, 'rb') as docx_file:
                        pdf_buffer = BytesIO(docx_file.read())
            
            # Limpiar archivos temporales
            try:
                os.unlink(temp_docx_path)
                os.unlink(temp_docx_out_path)
                # Cleanup: remove temp files if exist
                try:
                    if self.has_libreoffice and 'pdf_path' in locals() and pdf_path and os.path.exists(pdf_path):
                        os.unlink(pdf_path)
                except Exception:
                    pass
            except:
                pass
            
            pdf_buffer.seek(0)
            return pdf_buffer
            
        except Exception as e:
            logger.exception(f"Error generando PDF: {e}")
            raise

    def generate_entrega_telefono_pdf_by_asignacion(self, asignacion_id: int, coordinador_data: dict = None, empleado_data: dict = None, observaciones: list = None) -> BytesIO:
        """
        Realiza la consulta de la asignación y datos relacionados y genera el documento.

        - `asignacion_id`: id de la asignación a generar
        - `coordinador_data`: opcional, datos del coordinador (si no se provee, se usará 'Coordinador IT')
        - `empleado_data`: opcional, datos del empleado con identidad (si no se provee, se consulta por FK)
        """
        try:
            # Importar el servicio que consulta la BD
            from .service import DeviceService
        except Exception:
            raise RuntimeError('No se pudo importar DeviceService para consultar la asignación')

        svc = DeviceService()
        asignacion = svc.get_asignacion(asignacion_id)
        if not asignacion:
            raise FileNotFoundError(f"Asignación {asignacion_id} no encontrada")

        # Obtener dispositivo
        device_id = asignacion.get('fk_id_dispositivo')
        if not device_id:
            raise RuntimeError('Dispositivo no especificado en asignación')

        dispositivo = svc.get_device(device_id)
        if not dispositivo:
            raise FileNotFoundError(f"Dispositivo {device_id} no encontrado")

        # Obtener plan/línea
        plan_id = dispositivo.get('fk_id_plan')
        numero_linea = ''
        costo_plan = 0
        if plan_id:
            try:
                plan = svc.get_plane(plan_id)
                if plan:
                    numero_linea = plan.get('numero_linea', '')
                    costo_plan = plan.get('costo_plan', 0)
            except Exception:
                pass

        asignacion['numero_linea'] = numero_linea
        asignacion['costo_plan'] = costo_plan

        # Adjuntar componentes del dispositivo para que el generador pueda usar info de CPU/RAM/DISCO
        try:
            componentes = svc.list_components(device_id)
            # incluir en el dict del dispositivo para que generate_entrega_telefono_pdf los utilice
            dispositivo['componentes'] = componentes
        except Exception:
            dispositivo['componentes'] = []

        # Obtener empleado (usar empleado_data si se proporcionó, sino consultar por FK)
        if empleado_data is None:
            empleado_data = {'nombre_completo': ''}
            empleado_id = asignacion.get('fk_id_empleado')
            if empleado_id:
                try:
                    empleado = svc.get_empleado(empleado_id)
                    if empleado:
                        empleado_data = empleado
                except Exception:
                    pass

        # Coordinador
        if coordinador_data is None:
            coordinador_data = {'nombre': 'Coordinador IT'}

        # Llamar al generador existente
        # Intentar seleccionar plantilla adecuada según el tipo de dispositivo
        try:
            from admin_disp.services.docexp import select_template
            tpl_name = select_template(dispositivo.get('categoria') or '')
            # resolver ruta dentro de la carpeta `form` (misma ubicación que antes)
            tpl_full = os.path.join(os.path.dirname(__file__), '..', 'form', str(tpl_name))
        except Exception:
            tpl_full = None

        return self.generate_entrega_telefono_pdf(asignacion, dispositivo, empleado_data, coordinador_data, observaciones, template_path=tpl_full)
    
    # Replacement and template-mapping utilities moved to admin_disp.services.remapdata.RemapData
    
    def _convert_docx_to_pdf(self, docx_path: str) -> str:
        """
        Convierte un archivo DOCX a PDF usando LibreOffice.
        """
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        try:
            # Usar soffice (LibreOffice en Windows)
            subprocess.run([
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(docx_path),
                docx_path
            ], check=True, capture_output=True, timeout=30)
            
            logger.info(f"PDF generado exitosamente: {pdf_path}")
            return pdf_path
            
        except Exception as e:
            logger.error(f"Error convertiendo a PDF: {e}")
            raise RuntimeError(f"No se pudo convertir a PDF: {e}")
    
    @staticmethod
    def _format_date(date_obj) -> str:
        """Convierte un objeto datetime a formato DD/MM/YYYY"""
        if date_obj is None:
            return ""
        if hasattr(date_obj, 'strftime'):
            return date_obj.strftime('%d/%m/%Y')
        if isinstance(date_obj, str):
            try:
                date_parsed = datetime.fromisoformat(date_obj.replace('Z', '+00:00'))
                return date_parsed.strftime('%d/%m/%Y')
            except:
                return date_obj
        return str(date_obj)
